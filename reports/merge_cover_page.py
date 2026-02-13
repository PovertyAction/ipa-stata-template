#!/usr/bin/env python3
"""
Add a custom cover page to the generated report

This script removes the auto-generated title page from the Quarto output
and prepends a custom cover page.

Usage:
    python reports/merge_cover_page.py
    
This will:
1. Take the generated reports/analysis_report.docx
2. Remove the title/subtitle/author section
3. Prepend the cover page from reports/cover_page.docx
4. Save as reports/analysis_report_final.docx
"""

import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt


def remove_title_section(doc, num_elements=3):
    """
    Remove the first few elements (title, subtitle, author, date)
    
    Args:
        doc: Document object
        num_elements: Number of elements to remove from the beginning
    """
    # Remove first N paragraphs (usually title, subtitle, author, date, TOC title)
    for _ in range(min(num_elements, len(doc.paragraphs))):
        p = doc.paragraphs[0]
        p._element.getparent().remove(p._element)
    
    return doc


def prepend_cover_page(cover_path, report_path, output_path):
    """
    Prepend cover page to report
    
    Args:
        cover_path: Path to cover page DOCX
        report_path: Path to generated report DOCX (with title removed)
        output_path: Path for final output DOCX
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    
    print(f"📄 Loading cover page: {cover_path.name}")
    cover_doc = Document(str(cover_path))
    
    print(f"📊 Loading report: {report_path.name}")
    report_doc = Document(str(report_path))
    
    # Remove title section from report (first 4-5 paragraphs typically)
    print("🗑️  Removing auto-generated title section...")
    report_doc = remove_title_section(report_doc, num_elements=4)
    
    # Create new document starting with cover
    print("📑 Merging documents...")
    final_doc = Document(str(cover_path))
    
    # Add section break after cover page
    last_paragraph = final_doc.paragraphs[-1]
    run = last_paragraph.add_run()
    
    # Add page break
    run.add_break()
    
    # Add all content from report
    for element in report_doc.element.body:
        final_doc.element.body.append(element)
    
    # Save final document
    final_doc.save(str(output_path))
    print(f"✅ Final report saved: {output_path}")
    print(f"   Open with: {output_path}")


def main():
    # Set default paths
    script_dir = Path(__file__).parent
    project_root = script_dir.parent
    reports_dir = project_root / 'reports'
    
    cover_path = reports_dir / 'cover_page.docx'
    report_path = reports_dir / 'analysis_report.docx'
    output_path = reports_dir / 'analysis_report_final.docx'
    
    # Allow command line override
    if len(sys.argv) == 4:
        cover_path = Path(sys.argv[1])
        report_path = Path(sys.argv[2])
        output_path = Path(sys.argv[3])
    
    # Validate files
    if not cover_path.exists():
        print(f"❌ Error: Cover page not found: {cover_path}")
        print(f"\n💡 Create your cover page and save it as: {cover_path}")
        print("   The cover page should include:")
        print("   - Project title")
        print("   - Author names and affiliations")
        print("   - Date")
        print("   - Any institutional logos or branding")
        sys.exit(1)
    
    if not report_path.exists():
        print(f"❌ Error: Report not found: {report_path}")
        print(f"\n💡 Generate the report first:")
        print("   just render-docx reports/analysis_report.qmd")
        sys.exit(1)
    
    # Merge documents
    try:
        prepend_cover_page(cover_path, report_path, output_path)
    except Exception as e:
        print(f"❌ Error merging documents: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
